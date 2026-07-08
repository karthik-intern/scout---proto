import os
import tempfile
import gradio as gr
import fitz  # PyMuPDF
import pandas as pd
from docx import Document
from langchain_core.documents import Document as LCDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from sentence_transformers import SentenceTransformer
from langchain.embeddings.base import Embeddings
from langchain_chroma import Chroma
from groq import Groq

# ---------- Embedding Model ----------
_embedding_model = None

def get_embedding_model():
    global _embedding_model
    if _embedding_model is None:
        cache_dir = "/tmp/sentence_transformers"
        os.makedirs(cache_dir, exist_ok=True)
        _embedding_model = SentenceTransformer("all-MiniLM-L6-v2", cache_folder=cache_dir)
    return _embedding_model


class BGEEmbeddings(Embeddings):
    def __init__(self, model):
        self.model = model

    def embed_documents(self, texts):
        return self.model.encode(texts, normalize_embeddings=True).tolist()

    def embed_query(self, text):
        return self.model.encode(text, normalize_embeddings=True).tolist()


# ---------- Document Parsers ----------
def parse_pdf(file_path):
    documents = []
    pdf = fitz.open(file_path)
    for page_num in range(len(pdf)):
        page = pdf.load_page(page_num)
        text = page.get_text().strip()
        if text:
            documents.append(LCDocument(
                page_content=text,
                metadata={"source": os.path.basename(file_path), "page": page_num + 1}
            ))
    return documents


def parse_docx(file_path):
    doc = Document(file_path)
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return [LCDocument(page_content=text, metadata={"source": os.path.basename(file_path), "page": 1})]


def parse_excel(file_path):
    documents = []
    excel = pd.ExcelFile(file_path)
    for sheet in excel.sheet_names:
        df = pd.read_excel(file_path, sheet_name=sheet)
        text = df.astype(str).to_string(index=False)
        documents.append(LCDocument(
            page_content=text,
            metadata={"source": os.path.basename(file_path), "sheet": sheet}
        ))
    return documents


def parse_csv(file_path):
    df = pd.read_csv(file_path)
    text = df.astype(str).to_string(index=False)
    return [LCDocument(page_content=text, metadata={"source": os.path.basename(file_path), "page": 1})]


def parse_txt(file_path):
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    return [LCDocument(page_content=text, metadata={"source": os.path.basename(file_path), "page": 1})]


def parse_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return parse_pdf(file_path)
    elif ext == ".docx":
        return parse_docx(file_path)
    elif ext in [".xlsx", ".xls"]:
        return parse_excel(file_path)
    elif ext == ".csv":
        return parse_csv(file_path)
    elif ext == ".txt":
        return parse_txt(file_path)
    return []


# ---------- Global State ----------
vector_store = None
processed_files = []


# ---------- Process Documents ----------
def process_documents(files):
    global vector_store, processed_files

    if not files:
        return "⚠️ Please upload at least one document.", ""

    all_documents = []
    file_names = []
    status_lines = []

    for file in files:
        try:
            docs = parse_file(file.name)
            all_documents.extend(docs)
            name = os.path.basename(file.name)
            file_names.append(name)
            status_lines.append(f"✅ {name} → {len(docs)} section(s)")
        except Exception as e:
            status_lines.append(f"❌ {os.path.basename(file.name)} → Error: {e}")

    if not all_documents:
        return "❌ No text could be extracted from the uploaded files.", ""

    # Chunk
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000, chunk_overlap=200,
        separators=["\n\n", "\n", ". ", " ", ""]
    )
    chunks = []
    for doc in all_documents:
        if len(doc.page_content.strip()) > 50:
            chunks.extend(splitter.create_documents([doc.page_content], metadatas=[doc.metadata]))
        else:
            chunks.append(doc)

    # Build vector store
    model = get_embedding_model()
    embedding_fn = BGEEmbeddings(model)
    vector_store = Chroma.from_documents(
        documents=chunks,
        embedding=embedding_fn,
        collection_name="rag_docs"
    )

    processed_files.extend(file_names)
    status_lines.append(f"\n📦 {len(chunks)} chunks indexed into vector store")
    status_lines.append(f"🟢 Ready to answer questions!")

    files_list = "\n".join([f"• {f}" for f in processed_files])
    return "\n".join(status_lines), files_list


def add_documents(files):
    global vector_store, processed_files

    if not files:
        return "⚠️ Please upload files to add.", "\n".join([f"• {f}" for f in processed_files])

    if vector_store is None:
        return process_documents(files)

    all_documents = []
    file_names = []
    status_lines = []

    for file in files:
        try:
            docs = parse_file(file.name)
            all_documents.extend(docs)
            name = os.path.basename(file.name)
            file_names.append(name)
            status_lines.append(f"✅ {name} → {len(docs)} section(s) added")
        except Exception as e:
            status_lines.append(f"❌ {os.path.basename(file.name)} → {e}")

    if all_documents:
        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        chunks = []
        for doc in all_documents:
            if len(doc.page_content.strip()) > 50:
                chunks.extend(splitter.create_documents([doc.page_content], metadatas=[doc.metadata]))
            else:
                chunks.append(doc)

        model = get_embedding_model()
        vector_store.add_documents(chunks)
        processed_files.extend(file_names)
        status_lines.append(f"📦 {len(chunks)} new chunks added!")

    files_list = "\n".join([f"• {f}" for f in processed_files])
    return "\n".join(status_lines), files_list


# ---------- Chat ----------
def chat(message, history, groq_api_key):
    global vector_store

    if not groq_api_key or groq_api_key.strip() == "":
        groq_api_key = os.getenv("GROQ_API_KEY", "")

    if not groq_api_key:
        history.append((message, "⚠️ Please enter your Groq API key in the settings panel."))
        return "", history

    if vector_store is None:
        history.append((message, "📂 Please upload and process documents first before asking questions."))
        return "", history

    try:
        retriever = vector_store.as_retriever(search_type="similarity", search_kwargs={"k": 5})
        relevant_docs = retriever.invoke(message)

        context = "\n\n---\n\n".join([
            f"[Source: {doc.metadata.get('source', 'unknown')}, Page: {doc.metadata.get('page', 'N/A')}]\n{doc.page_content}"
            for doc in relevant_docs
        ])

        client = Groq(api_key=groq_api_key)
        response = client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are an intelligent document assistant. Answer questions based on the provided context. "
                        "Be concise, accurate, and cite the source document when possible. "
                        "If the answer is not in the context, say so clearly.\n\n"
                        f"Context:\n{context}"
                    )
                },
                {"role": "user", "content": message}
            ],
            temperature=0.3,
            max_tokens=1024
        )

        answer = response.choices[0].message.content

        sources = list(set([
            f"{doc.metadata.get('source', 'unknown')} (p.{doc.metadata.get('page', 'N/A')})"
            for doc in relevant_docs
        ]))
        sources_text = "\n\n📚 **Sources:** " + " | ".join(sources)

        history.append((message, answer + sources_text))

    except Exception as e:
        history.append((message, f"❌ Error: {str(e)}"))

    return "", history


def clear_all():
    global vector_store, processed_files
    vector_store = None
    processed_files = []
    return [], "", "", "🗑️ Knowledge base cleared."


# ---------- UI ----------
CSS = """
/* Overall theme */
.gradio-container {
    font-family: 'Inter', sans-serif !important;
    background: linear-gradient(135deg, #0f0c29, #302b63, #24243e) !important;
    min-height: 100vh;
}

/* Header */
.header-box {
    background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
    border-radius: 16px;
    padding: 24px 32px;
    margin-bottom: 20px;
    text-align: center;
    box-shadow: 0 8px 32px rgba(102, 126, 234, 0.4);
}

/* Panels */
.panel {
    background: rgba(255,255,255,0.05) !important;
    border: 1px solid rgba(255,255,255,0.1) !important;
    border-radius: 16px !important;
    backdrop-filter: blur(10px) !important;
}

/* Buttons */
.btn-primary {
    background: linear-gradient(90deg, #667eea, #764ba2) !important;
    border: none !important;
    border-radius: 10px !important;
    color: white !important;
    font-weight: 600 !important;
    transition: all 0.3s ease !important;
}

.btn-primary:hover {
    transform: translateY(-2px) !important;
    box-shadow: 0 8px 20px rgba(102, 126, 234, 0.5) !important;
}

.btn-secondary {
    background: rgba(255,255,255,0.1) !important;
    border: 1px solid rgba(255,255,255,0.2) !important;
    border-radius: 10px !important;
    color: white !important;
}

/* Chatbot */
.chatbot-container {
    border-radius: 16px !important;
    border: 1px solid rgba(255,255,255,0.1) !important;
}

/* Status boxes */
.status-box {
    background: rgba(16, 185, 129, 0.1) !important;
    border: 1px solid rgba(16, 185, 129, 0.3) !important;
    border-radius: 10px !important;
}

/* Input box */
input[type="text"], input[type="password"], textarea {
    background: rgba(255,255,255,0.08) !important;
    border: 1px solid rgba(255,255,255,0.15) !important;
    border-radius: 10px !important;
    color: white !important;
}

/* File upload */
.upload-area {
    border: 2px dashed rgba(102, 126, 234, 0.5) !important;
    border-radius: 12px !important;
    background: rgba(102, 126, 234, 0.05) !important;
}
"""

with gr.Blocks(
    css=CSS,
    title="RAG Document Assistant",
    theme=gr.themes.Soft(
        primary_hue="violet",
        secondary_hue="purple",
        neutral_hue="slate",
        font=gr.themes.GoogleFont("Inter")
    )
) as demo:

    # ---- Header ----
    gr.HTML("""
    <div class="header-box">
        <h1 style="color:white; font-size:2.2em; margin:0; font-weight:700;">
            📄 RAG Document Assistant
        </h1>
        <p style="color:rgba(255,255,255,0.85); margin:8px 0 0 0; font-size:1.1em;">
            Upload documents · Ask questions · Get AI-powered answers
        </p>
    </div>
    """)

    with gr.Row():
        # ---- Left Panel ----
        with gr.Column(scale=1, elem_classes="panel"):
            gr.HTML("<h3 style='color:#a78bfa; margin:0 0 12px 0;'>⚙️ Configuration</h3>")

            groq_key = gr.Textbox(
                label="🔑 Groq API Key",
                placeholder="Enter your Groq API key...",
                type="password",
                info="Get yours free at console.groq.com"
            )

            gr.HTML("<hr style='border-color:rgba(255,255,255,0.1); margin:16px 0;'>")
            gr.HTML("<h3 style='color:#a78bfa; margin:0 0 12px 0;'>📂 Upload Documents</h3>")

            file_upload = gr.File(
                label="Drop files here",
                file_count="multiple",
                file_types=[".pdf", ".docx", ".xlsx", ".csv", ".txt"],
                elem_classes="upload-area"
            )

            with gr.Row():
                process_btn = gr.Button("🚀 Process", variant="primary", elem_classes="btn-primary")
                add_btn = gr.Button("➕ Add More", elem_classes="btn-secondary")

            process_status = gr.Textbox(
                label="📊 Status",
                lines=6,
                interactive=False,
                elem_classes="status-box"
            )

            gr.HTML("<hr style='border-color:rgba(255,255,255,0.1); margin:16px 0;'>")
            gr.HTML("<h3 style='color:#a78bfa; margin:0 0 8px 0;'>📚 Knowledge Base</h3>")

            files_list = gr.Textbox(
                label="Indexed Documents",
                lines=4,
                interactive=False,
                placeholder="No documents loaded yet..."
            )

            clear_btn = gr.Button("🗑️ Clear All", variant="stop", elem_classes="btn-secondary")

            gr.HTML("""
            <div style='margin-top:16px; padding:12px; background:rgba(167,139,250,0.1);
                        border-radius:10px; border:1px solid rgba(167,139,250,0.2);'>
                <p style='color:rgba(255,255,255,0.6); font-size:0.8em; margin:0;'>
                    <b style='color:#a78bfa;'>Supported:</b> PDF · DOCX · XLSX · CSV · TXT<br>
                    <b style='color:#a78bfa;'>Powered by:</b> LangChain · ChromaDB · Groq LLaMA 3.1
                </p>
            </div>
            """)

        # ---- Right Panel (Chat) ----
        with gr.Column(scale=2, elem_classes="panel"):
            gr.HTML("<h3 style='color:#a78bfa; margin:0 0 12px 0;'>💬 Ask Questions</h3>")

            chatbot = gr.Chatbot(
                label="",
                height=520,
                show_label=False,
                avatar_images=("👤", "🤖"),
                bubble_full_width=False,
                elem_classes="chatbot-container"
            )

            with gr.Row():
                msg_input = gr.Textbox(
                    placeholder="Ask anything about your documents...",
                    label="",
                    show_label=False,
                    scale=5,
                    container=False
                )
                send_btn = gr.Button("Send ➤", variant="primary", scale=1, elem_classes="btn-primary")

            gr.HTML("""
            <div style='text-align:center; margin-top:8px;'>
                <span style='color:rgba(255,255,255,0.3); font-size:0.75em;'>
                    Press Enter or click Send · Powered by Groq LLaMA 3.1 8B
                </span>
            </div>
            """)

            with gr.Accordion("💡 Example Questions", open=False):
                gr.Examples(
                    examples=[
                        ["What is this document about?"],
                        ["Summarize the key points"],
                        ["What are the main conclusions?"],
                        ["List all important dates mentioned"],
                        ["Who are the key people mentioned?"],
                    ],
                    inputs=msg_input,
                    label=""
                )

    # ---- Events ----
    process_btn.click(
        fn=process_documents,
        inputs=[file_upload],
        outputs=[process_status, files_list]
    )

    add_btn.click(
        fn=add_documents,
        inputs=[file_upload],
        outputs=[process_status, files_list]
    )

    send_btn.click(
        fn=chat,
        inputs=[msg_input, chatbot, groq_key],
        outputs=[msg_input, chatbot]
    )

    msg_input.submit(
        fn=chat,
        inputs=[msg_input, chatbot, groq_key],
        outputs=[msg_input, chatbot]
    )

    clear_btn.click(
        fn=clear_all,
        outputs=[chatbot, files_list, process_status, process_status]
    )

if __name__ == "__main__":
    demo.launch(server_name="0.0.0.0", server_port=7860)
